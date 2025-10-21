import logging
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
import asyncpg
import asyncio
from docx import Document
from docx.shared import Cm

# from create_file import description, end_date

BOT_TOKEN = "8371430048:AAGWAVMuwGEvlZ4yTY2MAk0XnRag2RiZQf8"
DATABASE_URL = "postgresql://postgres:1234@db/subsidies"

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()


from docx import Document
from docx.shared import Cm
import subprocess

from docx import Document
from docx.shared import Cm
import subprocess
import os

class Create_file:

    def __init__(self, subsidy_data):
        self.name = subsidy_data.get('full_name', '')
        self.description = subsidy_data.get('description', '')
        self.link = subsidy_data.get('link', '')
        self.type_recipient = subsidy_data.get('type_recipient', '')
        self.admin_measure = subsidy_data.get('org', '')
        self.cofinancing = subsidy_data.get('cofinancing', '')
        self.start_date = subsidy_data.get('start_date', '')
        self.end_date = subsidy_data.get('end_date', '')
        self.requirements = subsidy_data.get('requirements', '')
        self.date = f"{self.start_date} - {self.end_date}" if self.start_date and self.end_date else ""

    def add_requirements_table(self, doc, text_with_reqs, header="Требования", style="Table Grid"):
        """Добавляет таблицу с требованиями"""
        items = [s.strip() for s in str(text_with_reqs).split("•")]
        items = [s for s in items if s]

        doc.add_heading(header, level=1)
        
        if not items:
            doc.add_paragraph("Требования не указаны")
            return None

        table = doc.add_table(rows=1, cols=1)
        table.style = style

        first_cell = table.rows[0].cells[0]
        first_cell.text = items[0]

        for s in items[1:]:
            row_cells = table.add_row().cells
            row_cells[0].text = s

        return table

    def add_table(self, doc, headers, rows, style="Table Grid", add_caption=None):
        """Добавляет таблицу в документ"""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = style
        hdr_cells = table.rows[0].cells
        for j, h in enumerate(headers):
            hdr_cells[j].text = str(h)
        
        for r in rows:
            row_cells = table.add_row().cells
            for j, val in enumerate(r):
                row_cells[j].text = str(val)
        
        if add_caption:
            doc.add_paragraph(add_caption)
        
        return table

    def build_docx(self, output_path="report.docx"):
        """Создает документ Word и возвращает путь к файлу"""
        doc = Document()
        
        # Настройка полей документа
        for section in doc.sections:
            section.left_margin = Cm(1.2)
            section.right_margin = Cm(1.2)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
        
        doc.add_heading("Отчёт по субсидии", level=0)

        # Основная информация
        headers1 = ["Наименование меры поддержки", "Краткое описание меры"]
        rows1 = [[self.name, self.description]]
        doc.add_heading("Основная информация", level=1)
        self.add_table(doc, headers1, rows1, style="Table Grid")

        doc.add_paragraph()

        # Детали меры поддержки
        headers1_1 = ["Якорь", "Позиция", "Отображение карточки", "Изображение", "Код продукта", "Уровни поддержки"]
        rows1_1 = [[" ", " ", "вкл", "нет", " ", "Региональная мера"]]
        doc.add_heading("Детали меры поддержки", level=1)
        self.add_table(doc, headers1_1, rows1_1, style="Table Grid")

        doc.add_paragraph()

        # Ссылки на подачу заявки
        headers2 = ["Сегмент", "Наименование кнопки", "Ссылка для перехода"]
        rows2 = [['MMB', "Подать заявку", self.link]]
        doc.add_heading("Подача заявки", level=1)
        self.add_table(doc, headers2, rows2, style="Table Grid")

        # Условия программы
        headers5 = ["Позиция", "Название условия", "Описание", "Иконка"]
        rows5 = [
            [1, 'Получатель поддержки', self.type_recipient, "free-icon-font-user-3917711.svg"],
            [2, "Администратор меры", self.admin_measure, "free-icon-font-bank-3914993.svg"],
            [3, "Софинансирование", self.cofinancing, "free-icon-font-badge-percent-7653146.svg"],
            [4, "Даты приема заявок", self.date, "free-icon-font-calendar-3917244.svg"],
            [5, "Механизм получения", "Подача заявки через портал", "free-icon-font-apps-3917618.svg"],
        ]
        doc.add_heading("Условия программы", level=1)
        self.add_table(doc, headers5, rows5, style="Table Grid")

        # Требования
        if self.requirements:
            self.add_requirements_table(doc, self.requirements)

        # Сохраняем документ
        doc.save(output_path)
        
        # Возвращаем путь к файлу, а не объект Document
        return output_path

    def convert_to_pdf(self, docx_path):
        """Конвертирует DOCX в PDF"""
        try:
            subprocess.run([
                "libreoffice", "--headless", "--convert-to", "pdf", 
                docx_path, "--outdir", os.path.dirname(docx_path)
            ], check=True, capture_output=True)
            
            pdf_path = docx_path.replace('.docx', '.pdf')
            return pdf_path if os.path.exists(pdf_path) else None
            
        except (subprocess.CalledProcessError, FileNotFoundError):
            return None
            
            # Состояния для FSM
class DocumentState(StatesGroup):
    waiting_for_code = State()

async def get_db_connection():
    return await asyncpg.connect(DATABASE_URL)

async def get_subsidies_from_db(min_days=15):
    connection = await get_db_connection()
    try:
        rows = await connection.fetch("""
            SELECT code, full_name, description, end_date
            FROM public.main_table 
            WHERE TO_DATE(SPLIT_PART(end_date, ' ', 1), 'DD.MM.YYYY') - CURRENT_DATE >= $1
            ORDER BY TO_DATE(SPLIT_PART(end_date, ' ', 1), 'DD.MM.YYYY') ASC
        """, min_days)
        
        return [dict(row) for row in rows]
    except Exception as e:
        logging.error(f"Database error: {e}")
        return []
    finally:
        await connection.close()

async def get_subsidy_by_code(code: str):
    """Получить субсидию по коду"""
    connection = await get_db_connection()
    try:
        row = await connection.fetchrow("""
            SELECT code, full_name, description, link, type_recipient, org, cofinancing, start_date, end_date, requirements
            FROM public.main_table 
            WHERE code = $1
        """, code)
        
        return dict(row) if row else None
    except Exception as e:
        logging.error(f"Database error: {e}")
        return None
    finally:
        await connection.close()

# async def generate_document(subsidy_data: dict) -> str:
#     """Генерация документа на основе данных субсидии"""
#     document = f"""
# 🏛️ **ДОКУМЕНТ ПО СУБСИДИИ**

# 📋 **Шифр субсидии:** {subsidy_data['code']}
# 🏷️ **Наименование:** {subsidy_data['full_name']}

# 📄 **ОПИСАНИЕ:**
# {subsidy_data['description']}

# 📅 **Срок отбора заявок до:** {subsidy_data['end_date']}
# """

#     if subsidy_data.get('amount'):
#         document += f"💵 **Размер финансирования:** {subsidy_data['amount']}\n\n"

#     if subsidy_data.get('requirements'):
#         document += f"📋 **ТРЕБОВАНИЯ К УЧАСТНИКАМ:**\n{subsidy_data['requirements']}\n\n"

#     if subsidy_data.get('organization'):
#         document += f"🏢 **Ответственная организация:** {subsidy_data['organization']}\n"

#     if subsidy_data.get('contact_info'):
#         document += f"📞 **Контактная информация:** {subsidy_data['contact_info']}\n"

#     document += f"\n---\n*Документ сгенерирован автоматически*"

#     return document

@dp.message(Command("start"))
async def cmd_start(message: types.Message):
    keyboard = InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="📊 Получить субсидии", callback_data="get_subsidies")],
            [InlineKeyboardButton(text="📄 Сформировать документ", callback_data="get_doc")]
        ]
    )
    
    await message.answer(
        "💼 **Бот информации о субсидиях**\n\n"
        "Выберите действие:",
        reply_markup=keyboard
    )

@dp.message(Command("get_doc"))
async def cmd_get_doc(message: types.Message, state: FSMContext):
    """Команда для формирования документа по коду субсидии"""
    await message.answer(
        "📄 **Формирование документа по субсидии**\n\n"
        "Пожалуйста, введите код субсидии:\n"
        "Пример: 25-961-25310-2-0176"
    )
    await state.set_state(DocumentState.waiting_for_code)

@dp.callback_query(lambda c: c.data == "get_doc")
async def process_get_doc(callback: types.CallbackQuery, state: FSMContext):
    """Обработчик кнопки формирования документа"""
    await callback.answer()
    await callback.message.answer(
        "📄 **Формирование документа по субсидии**\n\n"
        "Пожалуйста, введите код субсидии:\n"
        "Пример: 25-961-25310-2-0176"
    )
    await state.set_state(DocumentState.waiting_for_code)

@dp.message(DocumentState.waiting_for_code)
async def process_subsidy_code(message: types.Message, state: FSMContext):
    """Обработка введенного кода субсидии"""
    code = message.text.strip()
    
    # Проверяем формат кода (можно добавить более строгую проверку)
    if len(code) < 5:
        await message.answer("❌ Код слишком короткий. Пожалуйста, введите корректный код субсидии:")
        return
    
    await message.answer("🔍 Поиск субсидии в базе данных...")
    
    # Ищем субсидию по коду
    subsidy_data = await get_subsidy_by_code(code)
    
    if not subsidy_data:
        await message.answer(
            f"❌ Субсидия с кодом `{code}` не найдена.\n"
            "Проверьте правильность кода и попробуйте снова.\n\n"
            "Для повторной попытки введите код:"
        )
        return
    
    # Генерируем документ
    doc_generator = Create_file(subsidy_data)
    docx_path = f"subsidy_{code}.docx"
        
    # build_docx теперь возвращает путь к файлу, а не объект Document
    file_path = doc_generator.build_docx(docx_path)


    with open(file_path, 'rb') as file:
        await message.answer_document(
            document=types.BufferedInputFile(
                file.read(), 
                filename=f"subsidy_{code}.docx"
            ),
            caption=f"📄 Документ по субсидии {code}"
        )
    
    # Создаем клавиатуру для дополнительных действий
    keyboard = InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="📊 Все субсидии", callback_data="get_subsidies")],
            [InlineKeyboardButton(text="📄 Новый документ", callback_data="get_doc")]
        ]
    )
    
    await message.answer(document, reply_markup=keyboard)
    await state.clear()

@dp.callback_query(lambda c: c.data == "get_subsidies")
async def process_get_subsidies(callback: types.CallbackQuery):
    await callback.answer()
    
    # Показываем уведомление о загрузке
    await callback.message.edit_text("🔄 Загружаем актуальные данные...")
    
    subsidies = await get_subsidies_from_db()
    print(subsidies)
    
    if not subsidies:
        await callback.message.edit_text(
            "❌ На данный момент нет доступных субсидий.\n"
            "Попробуйте позже или обратитесь в поддержку."
        )
        return
    
    # Отправляем первую субсидию
    await send_subsidy_details(callback.message, subsidies, 0)

async def send_subsidy_details(message: types.Message, subsidies: list, index: int):
    subsidy = subsidies[index]
    
    keyboard = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(
                    text="⬅️ Назад", 
                    callback_data=f"subsidy_{max(0, index-1)}"
                ),
                InlineKeyboardButton(
                    text=f"{index+1}/{len(subsidies)}", 
                    callback_data="current"
                ),
                InlineKeyboardButton(
                    text="Вперёд ➡️", 
                    callback_data=f"subsidy_{min(len(subsidies)-1, index+1)}"
                )
            ],
            [
                InlineKeyboardButton(text="📋 Все субсидии", callback_data="all_subsidies"),
                InlineKeyboardButton(text="📄 Документ по коду", callback_data="get_doc")
            ]
        ]
    )
    
    text = (
        f"🏛️ **{subsidy['full_name']}**\n\n"
        f"💵 **Шифр:** {subsidy['code']}\n"
        f"📝 **Описание:** {subsidy['description']}\n"
        f"📅 **Отбор до:** {subsidy['end_date']}\n"
    )
    
    if subsidy.get('requirements'):
        text += f"📋 **Требования:** {subsidy['requirements']}\n"
    
    await message.edit_text(text, reply_markup=keyboard)

@dp.callback_query(lambda c: c.data.startswith('subsidy_'))
async def process_subsidy_navigation(callback: types.CallbackQuery):
    await callback.answer()
    
    # Получаем индекс из callback_data
    index = int(callback.data.split('_')[1])
    subsidies = await get_subsidies_from_db()
    
    if subsidies:
        await send_subsidy_details(callback.message, subsidies, index)

@dp.callback_query(lambda c: c.data == "all_subsidies")
async def process_all_subsidies(callback: types.CallbackQuery):
    await callback.answer()
    
    subsidies = await get_subsidies_from_db()
    
    if not subsidies:
        await callback.message.edit_text("❌ Нет доступных субсидий.")
        return
    
    text = "🏛️ **Все доступные субсидии:**\n\n"
    
    for i, subsidy in enumerate(subsidies, 1):
        # Используем существующие поля из вашей базы данных
        text += f"{i}. **{subsidy['full_name']}** - {subsidy['code']}\n"
    
    keyboard = InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="📋 Подробнее", callback_data="get_subsidies")],
            [InlineKeyboardButton(text="📄 Документ по коду", callback_data="get_doc")]
        ]
    )
    
    await callback.message.edit_text(text, reply_markup=keyboard)

# Обработчик для отмены ввода кода
@dp.message(Command("cancel"))
async def cmd_cancel(message: types.Message, state: FSMContext):
    current_state = await state.get_state()
    if current_state is None:
        return
    
    await state.clear()
    await message.answer(
        "❌ Ввод кода отменен.\n"
        "Для начала заново используйте /get_doc",
        reply_markup=InlineKeyboardMarkup(
            inline_keyboard=[[InlineKeyboardButton(text="📄 Сформировать документ", callback_data="get_doc")]]
        )
    )

async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())