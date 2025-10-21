# pip install python-docx
from docx import Document
from docx.shared import Inches
import pandas as pd
from docx.shared import Cm

df = pd.read_excel('downloads/Реестр отборов.xlsx')
idx = df.index[df["Шифр отбора"] == '25-934-25000-1-0080'] 
# print(idx)

name = df.loc[idx, "Полное наименование отбора"].squeeze()
description = df.loc[idx, "description"].squeeze()
link = df.loc[idx, "ссылка"].squeeze()
type_recipient = '<ul><li>' + df.loc[idx, "Тип получателей"].squeeze() + '</li></ul>'
admin_measure = '<ul><li>' + df.loc[idx, "Организация, предоставляющая субсидию"].squeeze() + '</li></ul>'
cofinancing = '<ul><li>' + df.loc[idx, "Софинансирование"].squeeze() + '</li></ul>'
start_date = df.loc[idx, "Дата начала приема заявок"].squeeze()
end_date = df.loc[idx, "Дата окончания приема заявок"].squeeze()
date = '<ul><li>' + start_date + " - " + end_date + '</li></ul>'
requirements = df.loc[idx, "requirements"].squeeze()
def add_requirements_table(doc, text_with_reqs, header="Требования", style="Table Grid"):
    # 1) Разбить по \n и убрать пустые
    items = [s.strip() for s in str(text_with_reqs).split("•")]
    items = [s for s in items if s]  # без пустых

    # 2) Создать таблицу: заголовок (опционально) и одна колонка
    doc.add_heading(header, level=1)
    table = doc.add_table(rows=1, cols=1)
    table.style = style

    # 3) Заполнить: каждая строка таблицы = одно требование
    # первая строка уже есть (row 0)
    first_cell = table.cell(0, 0)
    first_cell.text = items[0] if items else ""

    for s in items[1:]:
        row_cells = table.add_row().cells
        row_cells[0].text = s

    return table


def add_table(document, headers, rows, style="Table Grid", add_caption=None):
    # создаём таблицу с одной строкой под заголовки
    table = document.add_table(rows=1, cols=len(headers))
    table.style = style
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = str(h)
    # добавляем данные
    for r in rows:
        row_cells = table.add_row().cells
        for j, val in enumerate(r):
            row_cells[j].text = str(val)
    # при необходимости подпись (как абзац перед/после)
    
    if add_caption:
        p = document.add_paragraph(add_caption)
    return table

def build_docx(output_path="report.docx"):
    
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
    doc.add_heading("Отчёт с несколькими таблицами", level=0)

    # Таблица 1
    headers1 = ["Наименование меры поддержки", "Краткое описание меры"]
    rows1 = [
        [name, description]
    ]

    doc.add_heading("Карточка на сайт", level=1)
    add_table(doc, headers1, rows1, style="Table Grid")


    # Разделитель
    doc.add_paragraph()

    # Таблица 1
    headers1_1 = ["Якорь", "Позиция", "Отображение карточки программы", "Изображение", "Код продукта", "Уровни поддержки" ]
    rows1_1 = [
        [" ", " ", "вкл", "нет", " ", "Региональная мера"]
    ]

    doc.add_heading("Карточка на сайт", level=1)
    add_table(doc, headers1_1, rows1_1, style="Table Grid")


    # Разделитель
    doc.add_paragraph()

    # Таблица 2
    headers2 = ["Сегмент", "Наименование кнопки", "ссылка для перехода"]
    rows2 = [
        ['MMB', "Подать заявку", link],
       
    ]
    doc.add_heading("Кнопка целевого действия 1", level=1)
    add_table(doc, headers2, rows2, style="Table Grid")

    # Таблица 3 (динамическая ширина)
    headers3 = ["Сегмент", "Наименование кнопки", "ссылка для перехода"]
    rows3 = [
        ["MMB", "Подать заявку", link]
    ]
    doc.add_heading("Кнопка целевого действия 2", level=1)
    add_table(doc, headers3, rows3, style="Table Grid")
    

    headers4 = ["Наименование кнопки", "Ссылка для перехода"]
    rows4 = [
        ["Условие отбора", link]
    ]
    doc.add_heading("Дополнительная информация", level=1)
    add_table(doc, headers4, rows4, style="Table Grid")



    headers5 = ["Позиция", "Названия условия", "Детальное описание", "Путь к изображению", "Какая картинка должна быть"]
    rows5 = [
        [1,'Получатель поддержки', type_recipient, "free-icon-font-user-3917711.svg"],
        [2, "Администратор меры поддержки", admin_measure, "free-icon-font-bank-3914993.svg"],
        [3, "Софинансирование", cofinancing, "free-icon-font-badge-percent-7653146.svg"],
        [4, "Даты приема заявок", date, "free-icon-font-calendar-3917244.svg"],
        [5, "Механизм получения", "<ul><li>Для участия в конкурсе необходимо подать заявку на Портале предоставления мер финансовой государственной поддержки (promote.buget.gov.ru)</li></ul>", "free-icon-font-apps-3917618.svg"],

    ]
    
    doc.add_heading("Условия программы", level=1)
    add_table(doc, headers5, rows5, style="Table Grid")
  
    # headers6 = ["", "attribute_code", "website_attribute_name", "promt"]

    # rows6 = [
    #     ['Регион','Получатель поддержки',"", "free-icon-font-user-3917711.svg"],
    #     ["ЮЛ, ИП или ЮЛ и ИП", "Администратор меры поддержки", "","free-icon-font-bank-3914993.svg"],
    #     [3, "Софинансирование", "", "free-icon-font-badge-percent-7653146.svg"],
    #     [4, "Даты приема заявок", "", "free-icon-font-calendar-3917244.svg"],
    #     [5, "Механизм получения", "", "free-icon-font-apps-3917618.svg"],

    # ]
    
    # doc.add_heading("Таблица 1. Пользователи", level=1)
    # add_table(doc, headers6, rows6, style="Table Grid")

    add_requirements_table(doc, requirements)
  

    # Сохранение
    doc.save(output_path)

    import subprocess
    subprocess.run(["libreoffice", "--headless", "--convert-to", "pdf", output_path], check=True)


if __name__ == "__main__":
    build_docx("multi_tables.docx")
